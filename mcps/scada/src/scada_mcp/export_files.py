"""
File-based export (compatible with PHP ExportTools.php result shape: _type: file).
openpyxl + xhtml2pdf + python-docx -- server-side generation.
"""

from __future__ import annotations

import csv
import json
import logging
import os
import re
from collections import defaultdict
import time
from datetime import datetime
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Any

from . import db as dbmod
from .exports_http import (
    append_download_token,
    download_path,
    download_signed_path,
    public_download_signed_path,
    public_download_url,
)
from .log_value_cleanup import winsorize_tagvalue_rows
from .types import InstanceConfig

logger = logging.getLogger("scada_mcp.export_files")

_EXPORT_MAX_AGE_SEC = 86400
_LOG_FETCH_LIMIT = 50_000


def exports_dir(cfg: InstanceConfig) -> Path:
    d = cfg.base_dir / "exports"
    d.mkdir(parents=True, exist_ok=True)
    return d


def clean_old_exports(d: Path) -> None:
    if not d.is_dir():
        return
    cutoff = time.time() - _EXPORT_MAX_AGE_SEC
    for p in d.iterdir():
        try:
            if p.is_file() and p.stat().st_mtime < cutoff:
                p.unlink()
                logger.debug("Cleaned old export: %s", p.name)
        except OSError:
            pass


def safe_filename_part(name: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\-]", "_", name)[:120] or "export"


def _public_base_for_cfg(cfg: InstanceConfig | None) -> str | None:
    if cfg is None:
        return None
    u = getattr(cfg, "mcp_public_base_url", None)
    if isinstance(u, str) and u.strip():
        return u.strip().rstrip("/")
    return None


def file_ui_hints(filename: str, mime: str) -> dict[str, Any]:
    """
    Web / custom UI: fixed keys for icon and style (Excel, Word, PDF, ...).
    ui_kind: lowercase English; use with switch/case or CSS data-kind.
    """
    fn = (filename or "").strip()
    ext = ""
    if "." in fn:
        ext = fn.rsplit(".", 1)[-1].lower().strip()
    m = (mime or "").lower()

    by_ext: dict[str, tuple[str, str]] = {
        "xlsx": ("excel", "Excel"),
        "xls": ("excel", "Excel"),
        "csv": ("csv", "CSV"),
        "docx": ("word", "Word"),
        "doc": ("word", "Word"),
        "pdf": ("pdf", "PDF"),
        "json": ("json", "JSON"),
        "txt": ("text", "Text"),
        "html": ("html", "HTML"),
        "htm": ("html", "HTML"),
        "png": ("image", "Image"),
        "jpg": ("image", "Image"),
        "jpeg": ("image", "Image"),
        "gif": ("image", "Image"),
        "webp": ("image", "Image"),
        "zip": ("archive", "Archive"),
    }
    if ext in by_ext:
        uk, label = by_ext[ext]
        return {
            "file_extension": ext,
            "ui_kind": uk,
            "ui_label": label,
            "ui_icon_hint": uk,
        }

    if "spreadsheet" in m or m.endswith("sheet"):
        return {"file_extension": ext or "xlsx", "ui_kind": "excel", "ui_label": "Excel", "ui_icon_hint": "excel"}
    if m == "application/pdf" or "pdf" in m:
        return {"file_extension": ext or "pdf", "ui_kind": "pdf", "ui_label": "PDF", "ui_icon_hint": "pdf"}
    if "wordprocessingml" in m or m == "application/msword":
        return {"file_extension": ext or "docx", "ui_kind": "word", "ui_label": "Word", "ui_icon_hint": "word"}
    if "csv" in m or "comma-separated" in m:
        return {"file_extension": ext or "csv", "ui_kind": "csv", "ui_label": "CSV", "ui_icon_hint": "csv"}
    if "json" in m:
        return {"file_extension": ext or "json", "ui_kind": "json", "ui_label": "JSON", "ui_icon_hint": "json"}
    if m.startswith("text/"):
        return {"file_extension": ext or "txt", "ui_kind": "text", "ui_label": "Text", "ui_icon_hint": "text"}
    if m.startswith("image/"):
        return {"file_extension": ext, "ui_kind": "image", "ui_label": "Image", "ui_icon_hint": "image"}

    return {
        "file_extension": ext or "bin",
        "ui_kind": "other",
        "ui_label": "File",
        "ui_icon_hint": "file",
    }


def file_payload(
    path: Path,
    filename: str,
    mime: str,
    title: str,
    extra: dict[str, Any] | None = None,
    *,
    cfg: InstanceConfig | None = None,
) -> dict[str, Any]:
    kb = round(path.stat().st_size / 1024, 1) if path.exists() else 0.0
    out: dict[str, Any] = {
        "_type": "file",
        "title": title,
        "filename": filename,
        "filepath": str(path.resolve()),
        "mime_type": mime,
        "size_kb": kb,
    }
    out.update(file_ui_hints(filename, mime))
    if extra:
        out.update(extra)
    base = _public_base_for_cfg(cfg)
    if cfg is not None:
        out["instance_id"] = cfg.instance_id
        out["download_path"] = download_path(instance_id=cfg.instance_id, filename=filename)
        try:
            rel = path.resolve().relative_to(cfg.base_dir.resolve())
            out["exports_relative_path"] = str(rel).replace("\\", "/")
        except ValueError:
            out["exports_relative_path"] = f"exports/{filename}"
        if not base:
            out["download_hint"] = (
                "No download_url: set MCP_PUBLIC_BASE_URL or instance.yaml mcp_public_base_url (recommended). "
                "Use download_path or filepath for local access."
            )
    if cfg is not None:
        if base:
            du = public_download_url(
                public_base=base, instance_id=cfg.instance_id, filename=filename
            )
            out["download_url"] = du
            out["file_url"] = du
        else:
            out["download_url"] = out.get("download_path")
            out["file_url"] = out.get("download_path")

        if (
            cfg.auth
            and (cfg.auth.token_secret or "").strip()
            and (os.getenv("MCP_DOWNLOAD_SIGNED_URL") or "1").strip().lower() not in ("0", "false", "no")
        ):
            try:
                from .auth import TokenError, mint_file_download_token

                ttl = int(os.getenv("MCP_DOWNLOAD_TOKEN_TTL_SEC") or "3600")
                dtok = mint_file_download_token(
                    token_secret=cfg.auth.token_secret,
                    sub=cfg.instance_id,
                    filename=filename,
                    ttl_sec=ttl,
                )
                out["download_token_ttl_sec"] = max(60, ttl)
                if base:
                    signed = public_download_signed_path(
                        public_base=base, instance_id=cfg.instance_id, filename=filename, token=dtok
                    )
                    out["download_url"] = signed
                    out["file_url"] = signed
                    out["download_url_query"] = append_download_token(
                        public_download_url(
                            public_base=base, instance_id=cfg.instance_id, filename=filename
                        ),
                        dtok,
                    )
                else:
                    signed_rel = download_signed_path(
                        instance_id=cfg.instance_id, filename=filename, token=dtok
                    )
                    out["download_path_signed"] = signed_rel
                    out["download_url"] = signed_rel
                    out["file_url"] = signed_rel
            except TokenError:
                pass

        out.pop("download_hint", None)
        out["download_note"] = "download_url: /dt/ imzalı, Bearer gerekmez."
        out["presentation_rule_tr"] = (
            "URL'i PLAIN TEXT ver, markdown link [text](url) KULLANMA."
        )
    return out


def _winsorize_export_rows_by_logpid(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Multi logPId export: separate P1-P99 compression per parameter (sample = rows in that group)."""
    if not rows:
        return rows
    groups: dict[int, list[tuple[int, dict[str, Any]]]] = defaultdict(list)
    for i, r in enumerate(rows):
        pid = int(r.get("logPId") or 0)
        groups[pid].append((i, dict(r)))
    out: dict[int, dict[str, Any]] = {}
    for _pid, items in groups.items():
        chunk = [t[1] for t in items]
        clipped = winsorize_tagvalue_rows(chunk, value_key="tagValue")
        for (ix, _), wr in zip(items, clipped, strict=True):
            out[ix] = wr
    return [out[i] for i in range(len(rows))]


def _log_table_exists(cur: Any, node_id: int) -> bool:
    t = f"log_{int(node_id)}"
    cur.execute(
        "SELECT TABLE_NAME FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_SCHEMA='noktalog' AND TABLE_NAME=%s",
        (t,),
    )
    return cur.fetchone() is not None


def fetch_log_rows(
    cfg: InstanceConfig,
    node_id: int,
    start_date: str,
    end_date: str,
    log_param_ids: str,
) -> list[dict[str, Any]]:
    if not cfg.db:
        raise RuntimeError("DB config is missing for this instance.")
    nid = int(node_id)
    tname = f"log_{nid}"
    full = f"noktalog.`{tname}`"
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            if not _log_table_exists(cur, nid):
                return []
            where: list[str] = []
            params: list[Any] = []
            if (start_date or "").strip():
                where.append("l.logTime >= %s")
                params.append(start_date.strip())
            if (end_date or "").strip():
                where.append("l.logTime <= %s")
                params.append(end_date.strip())
            if (log_param_ids or "").strip():
                ids = [int(x) for x in log_param_ids.split(",") if x.strip().isdigit()]
                if ids:
                    ph = ",".join(["%s"] * len(ids))
                    where.append(f"l.logPId IN ({ph})")
                    params.extend(ids)
            wh = ("WHERE " + " AND ".join(where)) if where else ""
            sql = f"""
                SELECT l.logPId, lp.tagPath, lp.description,
                       l.tagValue, l.logTime
                FROM {full} l
                LEFT JOIN kbindb.logparameters lp ON l.logPId = lp.id
                {wh}
                ORDER BY l.logTime ASC
                LIMIT {_LOG_FETCH_LIMIT}
            """
            cur.execute(sql, tuple(params))
            return _winsorize_export_rows_by_logpid(list(cur.fetchall()))


def node_name(cfg: InstanceConfig, node_id: int) -> str:
    if not cfg.db:
        raise RuntimeError("DB config is missing for this instance.")
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT nName FROM kbindb.node WHERE id=%s", (int(node_id),))
            row = cur.fetchone()
    return str(row["nName"]) if row and row.get("nName") else f"Node {node_id}"


def build_xlsx(title: str, headers: list[str], rows: list[dict[str, Any]], filepath: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = (title[:31] if title else "Sheet1") or "Sheet1"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(fill_type="solid", start_color="1E3A5F", end_color="1E3A5F")
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(horizontal="center")
        c.border = border

    for ri, row in enumerate(rows, start=2):
        for ci, h in enumerate(headers, start=1):
            val = row.get(h) if isinstance(row, dict) else None
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.border = border

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 18

    wb.save(filepath)


def build_xlsx_simple(title: str, headers: list[str], table_rows: list[dict[str, Any]], filepath: Path) -> None:
    """table_rows: each row is a dict matching headers keys."""
    ordered = [{h: (r.get(h) if isinstance(r, dict) else "") for h in headers} for r in table_rows]
    build_xlsx(title, headers, ordered, filepath)


def build_csv(headers: list[str], rows: list[dict[str, Any]], filepath: Path) -> None:
    with filepath.open("w", encoding="utf-8-sig", newline="") as fp:
        w = csv.writer(fp, delimiter=";")
        w.writerow(headers)
        for r in rows:
            if isinstance(r, dict):
                w.writerow([r.get(h, "") for h in headers])
            else:
                w.writerow(list(r))


def build_csv_from_keys(
    headers: list[str], rows: list[dict[str, Any]], field_keys: list[str], filepath: Path
) -> None:
    """headers: visible headers; field_keys: row dict key order."""
    with filepath.open("w", encoding="utf-8-sig", newline="") as fp:
        w = csv.writer(fp, delimiter=";")
        w.writerow(headers)
        for r in rows:
            w.writerow([r.get(k, "") for k in field_keys])


def _rows_to_html_table(headers: list[str], rows: list[Any]) -> str:
    h = "<tr>" + "".join(f"<th>{escape(str(x))}</th>" for x in headers) + "</tr>"
    body = ""
    for row in rows:
        if isinstance(row, dict):
            vals = [row.get(h, "") for h in headers]
        else:
            vals = list(row)
        body += "<tr>" + "".join(f"<td>{escape(str(v))}</td>" for v in vals) + "</tr>"
    return f"<table>{h}{body}</table>"


def build_pdf_html(title: str, html_body: str, filepath: Path) -> None:
    from xhtml2pdf import pisa

    gen = datetime.now().strftime("%d.%m.%Y %H:%M")
    full = f"""<!DOCTYPE html><html><head><meta charset="UTF-8"/>
<style>
body {{ font-family: DejaVu Sans, Helvetica, sans-serif; font-size: 9pt; color: #1a1a1a; margin: 16px; }}
h1 {{ color: #1E3A5F; font-size: 14pt; border-bottom: 2px solid #1E3A5F; padding-bottom: 6px; }}
h2 {{ color: #2563eb; font-size: 11pt; margin-top: 12px; }}
table {{ width: 100%; border-collapse: collapse; margin: 8px 0; }}
th {{ background: #1E3A5F; color: #fff; padding: 4px 6px; text-align: left; font-size: 8pt; }}
td {{ padding: 4px 6px; border-bottom: 1px solid #ddd; font-size: 8pt; }}
tr:nth-child(even) td {{ background: #f0f4f8; }}
.footer {{ text-align: center; font-size: 7pt; color: #999; margin-top: 16px; border-top: 1px solid #ddd; padding-top: 6px; }}
</style></head><body>
<h1>{escape(title)}</h1>
{html_body}
<div class="footer">SCADA MCP -- Generated: {escape(gen)}</div>
</body></html>"""
    with filepath.open("wb") as out:
        status = pisa.CreatePDF(BytesIO(full.encode("utf-8")), dest=out, encoding="utf-8")
    err = getattr(status, "err", None)
    if err:
        raise RuntimeError("PDF generation failed (xhtml2pdf)")


def build_docx_report(title: str, sections: list[dict[str, Any]], filepath: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(title, 0)
    p0 = doc.add_paragraph()
    r0 = p0.add_run("Generated: " + datetime.now().strftime("%d.%m.%Y %H:%M"))
    r0.font.size = Pt(8)
    for sec in sections:
        if sec.get("heading"):
            doc.add_heading(str(sec["heading"]), level=2)
        if sec.get("text"):
            doc.add_paragraph(str(sec["text"]))
        tbl = sec.get("table")
        if isinstance(tbl, dict) and tbl.get("headers") and tbl.get("rows"):
            hdrs = tbl["headers"]
            table = doc.add_table(rows=1, cols=len(hdrs))
            table.style = "Table Grid"
            for i, hn in enumerate(hdrs):
                cell = table.rows[0].cells[i]
                cell.text = str(hn)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
            for row in tbl["rows"]:
                cells = table.add_row().cells
                vals = list(row.values()) if isinstance(row, dict) else list(row)
                for i, v in enumerate(vals):
                    if i < len(cells):
                        cells[i].text = str(v)
        doc.add_paragraph("")

    doc.save(filepath)


# --- Tool implementations ---


def export_log_data_impl(
    cfg: InstanceConfig,
    node_id: int,
    format: str = "xlsx",
    start_date: str = "",
    end_date: str = "",
    log_param_ids: str = "",
) -> Any:
    fmt = (format or "xlsx").lower()
    if fmt not in ("xlsx", "csv", "json"):
        return {"error": f"Invalid format: {format}. Supported: xlsx, csv, json"}

    nname = node_name(cfg, node_id)
    rows = fetch_log_rows(cfg, node_id, start_date, end_date, log_param_ids)
    if not rows:
        return {"error": "No data found for export."}

    clean_old_exports(exports_dir(cfg))
    d = exports_dir(cfg)
    stamp = datetime.now().strftime("%Y%m_%H%M%S")
    safe = safe_filename_part(nname)
    filename = f"log_{safe}_{stamp}.{fmt}"
    path = d / filename
    title = f"{nname} -- Log Data"

    headers = ["Time", "Parameter", "Description", "Value"]
    table_rows = [
        {
            "Time": r.get("logTime"),
            "Parameter": r.get("tagPath") or "",
            "Description": r.get("description") or "",
            "Value": r.get("tagValue"),
        }
        for r in rows
    ]

    if fmt == "xlsx":
        build_xlsx_simple(title, headers, table_rows, path)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif fmt == "csv":
        build_csv(headers, table_rows, path)
        mime = "text/csv"
    else:
        path.write_text(
            json.dumps(
                {
                    "node": {"id": node_id, "name": nname},
                    "date_range": {"start": start_date or "all", "end": end_date or "all"},
                    "record_count": len(rows),
                    "data": table_rows,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        mime = "application/json"

    return file_payload(
        path,
        filename,
        mime,
        title,
        {"record_count": len(rows), "format": fmt},
        cfg=cfg,
    )


def fetch_active_alarm_rows(cfg: InstanceConfig, limit: int) -> list[dict[str, Any]]:
    if not cfg.db:
        raise RuntimeError("DB config is missing for this instance.")
    lim = min(max(int(limit), 1), 500)
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            sql_state = """
                SELECT ast.pId, ap.name, ap.nid, n.nName AS node_name, ap.tagPath,
                       ap.minVal, ap.maxVal, ast.lastVal, ast.time AS alarm_time,
                       ap.alType, ap.alGroup, ap.comment
                FROM kbindb.alarmstate ast
                INNER JOIN kbindb.alarmparameters ap ON ast.pId = ap.pId
                LEFT JOIN kbindb.node n ON ap.nid = n.id
                WHERE ast.state = 1
                ORDER BY ast.time DESC
                LIMIT %s
            """
            try:
                cur.execute(sql_state, (lim,))
                rows = list(cur.fetchall())
                if rows:
                    return rows
            except Exception:
                pass
            sql_alt = """
                SELECT ap.pId, ap.name, ap.nid, n.nName AS node_name, ap.tagPath,
                       ap.minVal, ap.maxVal, ast.lastVal, ast.time AS alarm_time,
                       ap.alType, ap.alGroup, ap.comment, ast.alarmStateCurrent AS alarm_level
                FROM kbindb.alarmparameters ap
                LEFT JOIN kbindb.node n ON ap.nid = n.id
                INNER JOIN kbindb.alarmstate ast ON ap.id = ast.aid
                WHERE ast.alarmStateCurrent > 0
                ORDER BY ast.alarmStateCurrent DESC
                LIMIT %s
            """
            cur.execute(sql_alt, (lim,))
            return list(cur.fetchall())


def export_active_alarms_impl(cfg: InstanceConfig, format: str = "xlsx", limit: int = 500) -> Any:
    fmt = (format or "xlsx").lower()
    if fmt not in ("xlsx", "csv"):
        return {"error": "Invalid format. Supported: xlsx, csv"}

    rows = fetch_active_alarm_rows(cfg, limit)
    if not rows:
        return {"error": "No active alarm records found."}

    first = rows[0]
    keys = list(first.keys())
    header_labels = {
        "pId": "Alarm pId",
        "name": "Name",
        "nid": "Node ID",
        "node_name": "Station",
        "tagPath": "Tag",
        "minVal": "Min",
        "maxVal": "Max",
        "lastVal": "Last Value",
        "alarm_time": "Alarm Time",
        "alType": "Type",
        "alGroup": "Group",
        "comment": "Note",
        "alarm_level": "Alarm Level",
    }
    headers = [header_labels.get(k, k) for k in keys]

    table_rows: list[dict[str, Any]] = []
    for r in rows:
        line: dict[str, Any] = {}
        for k in keys:
            v = r.get(k)
            line[k] = "" if v is None else v
        table_rows.append(line)

    clean_old_exports(exports_dir(cfg))
    d = exports_dir(cfg)
    stamp = datetime.now().strftime("%Y%m_%H%M%S")
    ext = "xlsx" if fmt == "xlsx" else "csv"
    filename = f"active_alarms_{stamp}.{ext}"
    path = d / filename
    title = "Active Alarms -- " + datetime.now().strftime("%d.%m.%Y %H:%M")

    if fmt == "xlsx":
        excel_rows = []
        for r in table_rows:
            excel_rows.append({headers[i]: r[keys[i]] for i in range(len(keys))})
        build_xlsx_simple(title, headers, excel_rows, path)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        build_csv_from_keys(headers, table_rows, keys, path)
        mime = "text/csv"

    return file_payload(path, filename, mime, title, {"record_count": len(rows), "format": fmt}, cfg=cfg)


def _build_and_payload(
    *,
    cfg: InstanceConfig,
    rows: list[dict[str, Any]],
    header_labels: dict[str, str],
    format: str,
    filename_base: str,
    title_prefix: str,
    extra: dict[str, Any] | None = None,
) -> Any:
    """Tek tip export yardimcisi: rows + header_labels -> xlsx/csv + file_payload."""
    fmt = (format or "xlsx").lower()
    if fmt not in ("xlsx", "csv"):
        return {"error": "Invalid format. Supported: xlsx, csv"}
    if not rows:
        return {"error": "No records found."}

    keys = list(rows[0].keys())
    headers = [header_labels.get(k, k) for k in keys]

    table_rows: list[dict[str, Any]] = []
    for r in rows:
        line: dict[str, Any] = {}
        for k in keys:
            v = r.get(k)
            line[k] = "" if v is None else v
        table_rows.append(line)

    clean_old_exports(exports_dir(cfg))
    d = exports_dir(cfg)
    stamp = datetime.now().strftime("%Y%m_%H%M%S")
    filename = f"{filename_base}_{stamp}.{fmt}"
    path = d / filename
    title = f"{title_prefix} -- " + datetime.now().strftime("%d.%m.%Y %H:%M")

    if fmt == "xlsx":
        excel_rows = []
        for r in table_rows:
            excel_rows.append({headers[i]: r[keys[i]] for i in range(len(keys))})
        build_xlsx_simple(title, headers, excel_rows, path)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        build_csv_from_keys(headers, table_rows, keys, path)
        mime = "text/csv"

    meta = {"record_count": len(rows), "format": fmt}
    if extra:
        meta.update(extra)
    return file_payload(path, filename, mime, title, meta, cfg=cfg)


def export_nodes_impl(
    cfg: InstanceConfig,
    format: str = "xlsx",
    nType: str = "",
    keyword: str = "",
    only_active: bool = False,
) -> Any:
    """Tum node'lari tek cagrida Excel/CSV. Filtreler: nType, keyword (nName LIKE), only_active."""
    if not cfg.db:
        return {"error": "DB config missing"}
    where: list[str] = []
    params: list[Any] = []
    nt = (nType or "").strip()
    if nt and nt.lstrip("-").isdigit():
        where.append("n.nType = %s")
        params.append(int(nt))
    if (keyword or "").strip():
        where.append("n.nName LIKE %s")
        params.append(f"%{keyword.strip()}%")
    if only_active:
        where.append("n.nState >= 0")
    wh = ("WHERE " + " AND ".join(where)) if where else ""
    sql = f"""
        SELECT n.id, n.nName, n.nView, n.nType, n.nState, n.nPath, n.nBase,
               pt.name AS urun_tipi, pt.category AS urun_kategori
        FROM node n
        LEFT JOIN node_product_type pt ON n.nView = pt.nView
        {wh}
        ORDER BY n.id ASC
    """
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute(sql, tuple(params))
            rows = list(cur.fetchall())
    header_labels = {
        "id": "Node ID",
        "nName": "Ad",
        "nView": "Ekran Tipi (nView)",
        "nType": "Tip Kodu",
        "nState": "Durum",
        "nPath": "Yol",
        "nBase": "Parent ID",
        "urun_tipi": "Urun Tipi",
        "urun_kategori": "Kategori",
    }
    return _build_and_payload(
        cfg=cfg, rows=rows, header_labels=header_labels, format=format,
        filename_base="nodelar", title_prefix="Nodelar",
        extra={"filter_nType": nt, "filter_keyword": keyword, "only_active": only_active},
    )


def export_alarm_definitions_impl(
    cfg: InstanceConfig,
    format: str = "xlsx",
    nodeId: int = 0,
) -> Any:
    """alarmparameters: tum alarm TANIMLARI. Aktif alarm icin export_alarms_history kullan."""
    if not cfg.db:
        return {"error": "DB config missing"}
    where = ""
    params: list[Any] = []
    if int(nodeId or 0) > 0:
        where = "WHERE ap.nid = %s"
        params.append(int(nodeId))
    sql = f"""
        SELECT ap.pId, ap.name, ap.nid, n.nName AS node_adi, ap.tagPath,
               ap.minVal, ap.maxVal, ap.alType, ap.alGroup, ap.alGroupPath, ap.comment
        FROM alarmparameters ap
        LEFT JOIN node n ON ap.nid = n.id
        {where}
        ORDER BY ap.nid ASC, ap.pId ASC
    """
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute(sql, tuple(params))
            rows = list(cur.fetchall())
    header_labels = {
        "pId": "pId", "name": "Ad", "nid": "Node ID", "node_adi": "Node Adi",
        "tagPath": "Tag", "minVal": "Min", "maxVal": "Max",
        "alType": "Tip", "alGroup": "Grup", "alGroupPath": "Grup Yolu", "comment": "Aciklama",
    }
    return _build_and_payload(
        cfg=cfg, rows=rows, header_labels=header_labels, format=format,
        filename_base="alarm_tanimlari", title_prefix="Alarm Tanimlari",
    )


def export_alarms_history_impl(
    cfg: InstanceConfig,
    format: str = "xlsx",
    start_date: str = "",
    end_date: str = "",
    only_active: bool = False,
    limit: int = 5000,
) -> Any:
    """alarmstate zaman araligi export. Gercek history degil — alarmstate.time'a gore."""
    if not cfg.db:
        return {"error": "DB config missing"}
    lim = min(max(int(limit or 5000), 1), 20000)
    where: list[str] = []
    params: list[Any] = []
    if only_active:
        where.append("ast.state = 1")
    if (start_date or "").strip():
        where.append("ast.time >= %s")
        params.append(start_date.strip())
    if (end_date or "").strip():
        where.append("ast.time <= %s")
        params.append(end_date.strip())
    wh = ("WHERE " + " AND ".join(where)) if where else ""
    sql = f"""
        SELECT ast.pId, ap.name, ap.nid, n.nName AS node_adi, ap.tagPath,
               ap.minVal, ap.maxVal, ast.lastVal, ast.state,
               ast.time AS alarm_zamani, ap.alType, ap.alGroup, ap.comment
        FROM alarmstate ast
        INNER JOIN alarmparameters ap ON ast.pId = ap.pId
        LEFT JOIN node n ON ap.nid = n.id
        {wh}
        ORDER BY ast.time DESC
        LIMIT %s
    """
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute(sql, (*params, lim))
            rows = list(cur.fetchall())
    header_labels = {
        "pId": "pId", "name": "Ad", "nid": "Node ID", "node_adi": "Node Adi",
        "tagPath": "Tag", "minVal": "Min", "maxVal": "Max", "lastVal": "Son Deger",
        "state": "Aktif (1=Evet)", "alarm_zamani": "Zaman",
        "alType": "Tip", "alGroup": "Grup", "comment": "Aciklama",
    }
    return _build_and_payload(
        cfg=cfg, rows=rows, header_labels=header_labels, format=format,
        filename_base="alarm_gecmisi", title_prefix="Alarm Gecmisi",
        extra={"only_active": only_active, "start_date": start_date, "end_date": end_date},
    )


def export_user_groups_impl(cfg: InstanceConfig, format: str = "xlsx") -> Any:
    """Tum kullanici gruplari."""
    if not cfg.db:
        return {"error": "DB config missing"}
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute(
                """
                SELECT g.id, g.gName, g.gParent,
                       COUNT(u.id) AS kullanici_sayisi
                FROM user_groups g
                LEFT JOIN users u ON u.gid = g.id
                GROUP BY g.id, g.gName, g.gParent
                ORDER BY g.gName ASC
                """
            )
            rows = list(cur.fetchall())
    header_labels = {
        "id": "Grup ID", "gName": "Ad", "gParent": "Parent ID",
        "kullanici_sayisi": "Kullanici Sayisi",
    }
    return _build_and_payload(
        cfg=cfg, rows=rows, header_labels=header_labels, format=format,
        filename_base="kullanici_gruplari", title_prefix="Kullanici Gruplari",
    )


def export_users_impl(
    cfg: InstanceConfig,
    format: str = "xlsx",
    status: str = "all",
    company: str = "",
    city: str = "",
) -> Any:
    """Tum kullanicilari tek seferde Excel/CSV olarak export eder.
    LLM uzerinden data gecirmez — DB'den direkt okur.
    """
    fmt = (format or "xlsx").lower()
    if fmt not in ("xlsx", "csv"):
        return {"error": "Invalid format. Supported: xlsx, csv"}
    if not cfg.db:
        return {"error": "DB config missing"}

    where: list[str] = []
    params: list[Any] = []
    st = (status or "all").lower()
    if st == "active":
        where.append("u.uEnable = 1")
    elif st in ("inactive", "passive", "pasif"):
        where.append("u.uEnable = 0")
    if (company or "").strip():
        where.append("u.com_info LIKE %s")
        params.append(f"%{company.strip()}%")
    if (city or "").strip():
        where.append("u.uCity LIKE %s")
        params.append(f"%{city.strip()}%")
    wh = ("WHERE " + " AND ".join(where)) if where else ""
    sql = f"""
        SELECT u.id, u.uFirstName, u.uLastName, u.uName, u.uTel, u.uMail,
               u.uTitle, u.com_info, u.uCity, u.uLevel, u.uType,
               g.gName AS grup_adi,
               CASE WHEN u.uEnable = 1 THEN 'Aktif' ELSE 'Pasif' END AS durum,
               u.lastLogin, u.createdTime
        FROM users u LEFT JOIN user_groups g ON u.gid = g.id
        {wh} ORDER BY u.id DESC
    """
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute(sql, tuple(params))
            rows = list(cur.fetchall())

    if not rows:
        return {"error": "No users found for given filters."}

    header_labels = {
        "id": "ID",
        "uFirstName": "Ad",
        "uLastName": "Soyad",
        "uName": "Kullanici Adi",
        "uTel": "Telefon",
        "uMail": "Email",
        "uTitle": "Unvan",
        "com_info": "Sirket",
        "uCity": "Sehir",
        "uLevel": "Yetki Seviyesi",
        "uType": "Tip",
        "grup_adi": "Grup",
        "durum": "Durum",
        "lastLogin": "Son Giris",
        "createdTime": "Kayit Tarihi",
    }
    keys = list(rows[0].keys())
    headers = [header_labels.get(k, k) for k in keys]

    table_rows: list[dict[str, Any]] = []
    for r in rows:
        line: dict[str, Any] = {}
        for k in keys:
            v = r.get(k)
            line[k] = "" if v is None else v
        table_rows.append(line)

    clean_old_exports(exports_dir(cfg))
    d = exports_dir(cfg)
    stamp = datetime.now().strftime("%Y%m_%H%M%S")
    filename = f"kullanicilar_{stamp}.{fmt}"
    path = d / filename
    title = f"Kullanicilar ({st}) -- " + datetime.now().strftime("%d.%m.%Y %H:%M")

    if fmt == "xlsx":
        excel_rows = []
        for r in table_rows:
            excel_rows.append({headers[i]: r[keys[i]] for i in range(len(keys))})
        build_xlsx_simple(title, headers, excel_rows, path)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        build_csv_from_keys(headers, table_rows, keys, path)
        mime = "text/csv"

    return file_payload(
        path, filename, mime, title,
        {"record_count": len(rows), "format": fmt, "status": st},
        cfg=cfg,
    )


def export_custom_data_impl(
    cfg: InstanceConfig,
    title: str,
    headers: str,
    rows_json: str,
    format: str = "xlsx",
) -> Any:
    fmt = (format or "xlsx").lower()
    if fmt not in ("xlsx", "csv", "json", "pdf"):
        return {"error": "Invalid format. Supported: xlsx, csv, json, pdf"}

    header_list = [x.strip() for x in (headers or "").split(",") if x.strip()]
    try:
        rows = json.loads(rows_json or "[]")
    except json.JSONDecodeError as e:
        return {"error": f"rowsJson invalid JSON: {e}"}
    if not isinstance(rows, list) or not rows:
        return {"error": "rowsJson must be a valid JSON array."}

    if isinstance(rows[0], dict) and header_list:
        row_keys = list(rows[0].keys())
        missing = [h for h in header_list if h not in rows[0]]
        if missing:
            return {
                "error": (
                    "Column names in headers must exactly match the keys in rowsJson objects."
                ),
                "headers_provided": header_list,
                "row_keys_example": row_keys,
                "fix_example": {
                    "headers": ",".join(row_keys),
                    "description": "Make headers value match this comma-separated list, or match JSON keys to headers.",
                },
            }

    clean_old_exports(exports_dir(cfg))
    d = exports_dir(cfg)
    safe_t = safe_filename_part(title)
    stamp = datetime.now().strftime("%Y%m_%H%M%S")
    filename = f"export_{safe_t}_{stamp}.{fmt}"
    path = d / filename

    if fmt == "xlsx":
        norm = []
        for r in rows:
            if isinstance(r, dict):
                norm.append({h: r.get(h, "") for h in header_list})
            else:
                lv = list(r) if isinstance(r, (list, tuple)) else [r]
                norm.append({header_list[i]: lv[i] if i < len(lv) else "" for i in range(len(header_list))})
        build_xlsx_simple(title or "Export", header_list, norm, path)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif fmt == "csv":
        norm = []
        for r in rows:
            if isinstance(r, dict):
                norm.append({h: r.get(h, "") for h in header_list})
            else:
                lv = list(r) if isinstance(r, (list, tuple)) else [r]
                norm.append({header_list[i]: lv[i] if i < len(lv) else "" for i in range(len(header_list))})
        build_csv(header_list, norm, path)
        mime = "text/csv"
    elif fmt == "json":
        path.write_text(
            json.dumps(
                {"title": title, "columns": header_list, "record_count": len(rows), "data": rows},
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        mime = "application/json"
    else:
        body = _rows_to_html_table(header_list, rows)
        build_pdf_html(title or "Export", body, path)
        mime = "application/pdf"

    return file_payload(
        path, filename, mime, title or "Export", {"record_count": len(rows), "format": fmt}, cfg=cfg
    )


def generate_scada_report_impl(cfg: InstanceConfig, format: str = "pdf") -> Any:
    if not cfg.db:
        raise RuntimeError("DB config is missing for this instance.")
    fmt = (format or "pdf").lower()
    if fmt not in ("pdf", "docx"):
        return {"error": f"Invalid format: {format}. Supported: pdf, docx"}

    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT COUNT(*) cnt, nType FROM kbindb.node GROUP BY nType ORDER BY cnt DESC")
            node_stats = list(cur.fetchall())
            total_nodes = sum(int(r["cnt"]) for r in node_stats)

            cur.execute("SELECT COUNT(*) AS c FROM kbindb.users")
            user_count = int(cur.fetchone()["c"])

            cur.execute("SELECT COUNT(*) AS c FROM kbindb.alarmparameters")
            alarm_count = int(cur.fetchone()["c"])

            cur.execute("SELECT COUNT(*) AS c FROM kbindb.alarmstate WHERE alarmStateCurrent > 0")
            active_alarms = int(cur.fetchone()["c"])

            cur.execute("SELECT COUNT(*) AS c FROM kbindb._tagoku")
            tag_count = int(cur.fetchone()["c"])

            cur.execute("SELECT COUNT(*) AS c FROM kbindb.logparameters")
            log_param_count = int(cur.fetchone()["c"])

            cd_count = 0
            try:
                cur.execute("SELECT COUNT(*) AS c FROM dbdataexchanger.channeldevice")
                cd_count = int(cur.fetchone()["c"])
            except Exception:
                cd_count = 0

            top_alarms: list[dict[str, Any]] = []
            try:
                cur.execute(
                    """
                    SELECT ap.id, n.nName, ap.tagPath, ap.description, ast.alarmStateCurrent
                    FROM kbindb.alarmparameters ap
                    LEFT JOIN kbindb.node n ON ap.nid = n.id
                    LEFT JOIN kbindb.alarmstate ast ON ap.id = ast.aid
                    WHERE ast.alarmStateCurrent > 0
                    ORDER BY ast.alarmStateCurrent DESC
                    LIMIT 20
                    """
                )
                top_alarms = list(cur.fetchall())
            except Exception:
                top_alarms = []

    clean_old_exports(exports_dir(cfg))
    d = exports_dir(cfg)
    stamp = datetime.now().strftime("%Y%m_%H%M%S")
    title = "SCADA System Report -- " + datetime.now().strftime("%d.%m.%Y")

    if fmt == "pdf":
        stats_html = (
            "<table><tr>"
            + f"<td><strong>Total Nodes</strong><br/>{total_nodes}</td>"
            + f"<td><strong>Users</strong><br/>{user_count}</td>"
            + f"<td><strong>Live Tags</strong><br/>{tag_count}</td>"
            + f"<td><strong>Alarm Params</strong><br/>{alarm_count}</td>"
            + f"<td><strong>Active Alarms</strong><br/>{active_alarms}</td>"
            + f"<td><strong>Log Params</strong><br/>{log_param_count}</td>"
            + f"<td><strong>Channel Devices</strong><br/>{cd_count}</td>"
            + "</tr></table>"
        )
        nt = "<h2>Node Distribution</h2><table><tr><th>Node Type</th><th>Count</th></tr>"
        for ns_row in node_stats:
            nt += f"<tr><td>{escape(str(ns_row.get('nType')))}</td><td>{ns_row.get('cnt')}</td></tr>"
        nt += "</table>"

        ah = "<h2>Active Alarms (Top 20)</h2>"
        if not top_alarms:
            ah += "<p>No active alarms found.</p>"
        else:
            ah += "<table><tr><th>Node</th><th>Parameter</th><th>Description</th><th>Status</th></tr>"
            for a in top_alarms:
                ah += (
                    f"<tr><td>{escape(str(a.get('nName') or ''))}</td>"
                    f"<td>{escape(str(a.get('tagPath') or ''))}</td>"
                    f"<td>{escape(str(a.get('description') or ''))}</td>"
                    f"<td>{a.get('alarmStateCurrent')}</td></tr>"
                )
            ah += "</table>"

        filename = f"scada_report_{stamp}.pdf"
        path = d / filename
        build_pdf_html(title, stats_html + nt + ah, path)
        return file_payload(
            path,
            filename,
            "application/pdf",
            title,
            {"content": "SCADA system summary report", "node_count": total_nodes, "active_alarms": active_alarms},
            cfg=cfg,
        )

    # docx
    sections: list[dict[str, Any]] = [
        {
            "heading": "General Statistics",
            "table": {
                "headers": ["Metric", "Value"],
                "rows": [
                    ["Total Nodes", total_nodes],
                    ["Users", user_count],
                    ["Live Tags", tag_count],
                    ["Alarm Parameters", alarm_count],
                    ["Active Alarms", active_alarms],
                    ["Log Parameters", log_param_count],
                    ["Channel Devices", cd_count],
                ],
            },
        },
        {
            "heading": "Node Distribution",
            "table": {
                "headers": ["Node Type", "Count"],
                "rows": [[ns_row.get("nType"), ns_row.get("cnt")] for ns_row in node_stats],
            },
        },
    ]
    if top_alarms:
        sections.append(
            {
                "heading": "Active Alarms (Top 20)",
                "table": {
                    "headers": ["Node", "Parameter", "Description", "Status"],
                    "rows": [
                        [a.get("nName"), a.get("tagPath"), a.get("description"), a.get("alarmStateCurrent")]
                        for a in top_alarms
                    ],
                },
            }
        )
    else:
        sections.append({"heading": "Active Alarms", "text": "No active alarms found."})

    filename = f"scada_report_{stamp}.docx"
    path = d / filename
    build_docx_report(title, sections, path)
    return file_payload(
        path,
        filename,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        title,
        {"content": "SCADA system summary report", "node_count": total_nodes, "active_alarms": active_alarms},
        cfg=cfg,
    )


def generate_node_report_impl(cfg: InstanceConfig, node_id: int, format: str = "pdf") -> Any:
    if not cfg.db:
        raise RuntimeError("DB config is missing for this instance.")
    fmt = (format or "pdf").lower()
    if fmt not in ("pdf", "docx"):
        return {"error": "Invalid format. Supported: pdf, docx"}

    nid = int(node_id)
    with dbmod.connect(cfg.db) as conn:
        with conn.cursor() as cur:
            cur.execute("SELECT * FROM kbindb.node WHERE id=%s", (nid,))
            node = cur.fetchone()
            if not node:
                return {"error": f"Node {nid} not found."}

            cur.execute(
                "SELECT tagPath, tagValue, readTime FROM kbindb._tagoku WHERE nid=%s ORDER BY tagPath",
                (nid,),
            )
            tag_rows = list(cur.fetchall())

            cur.execute(
                """
                SELECT ap.tagPath, ap.description, ast.alarmStateCurrent
                FROM kbindb.alarmparameters ap
                LEFT JOIN kbindb.alarmstate ast ON ap.id = ast.aid
                WHERE ap.nid=%s
                ORDER BY ast.alarmStateCurrent DESC
                """,
                (nid,),
            )
            alarm_rows = list(cur.fetchall())

            cur.execute(
                """
                SELECT id, tagPath, description, rangeMin, rangeMax, logInterval
                FROM kbindb.logparameters WHERE nid=%s ORDER BY tagPath
                """,
                (nid,),
            )
            log_param_rows = list(cur.fetchall())

    node_name_s = str(node.get("nName") or f"Node {nid}")
    title = f"{node_name_s} -- Detail Report -- " + datetime.now().strftime("%d.%m.%Y")

    clean_old_exports(exports_dir(cfg))
    d = exports_dir(cfg)
    stamp = datetime.now().strftime("%Y%m_%H%M%S")
    safe = safe_filename_part(node_name_s)

    if fmt == "pdf":
        html = (
            f"<p><strong>Node ID:</strong> {nid} | <strong>Type:</strong> {escape(str(node.get('nType')))} | "
            f"<strong>Path:</strong> {escape(str(node.get('nPath') or ''))}</p>"
        )
        html += "<h2>Live Tag Values</h2><table><tr><th>Tag</th><th>Value</th><th>Last Read</th></tr>"
        for t in tag_rows:
            html += (
                f"<tr><td>{escape(str(t.get('tagPath')))}</td>"
                f"<td>{escape(str(t.get('tagValue')))}</td>"
                f"<td>{escape(str(t.get('readTime')))}</td></tr>"
            )
        html += "</table>"

        active_count = sum(1 for a in alarm_rows if int(a.get("alarmStateCurrent") or 0) > 0)
        html += f"<h2>Alarm Parameters ({active_count} active / {len(alarm_rows)} total)</h2>"
        if alarm_rows:
            html += "<table><tr><th>Parameter</th><th>Description</th><th>Status</th></tr>"
            for a in alarm_rows:
                st = (
                    '<b style="color:red">ACTIVE</b>'
                    if int(a.get("alarmStateCurrent") or 0) > 0
                    else "Normal"
                )
                html += (
                    f"<tr><td>{escape(str(a.get('tagPath')))}</td>"
                    f"<td>{escape(str(a.get('description')))}</td><td>{st}</td></tr>"
                )
            html += "</table>"

        html += f"<h2>Log Parameters ({len(log_param_rows)})</h2>"
        if log_param_rows:
            html += "<table><tr><th>Tag</th><th>Description</th><th>Min</th><th>Max</th><th>Log Interval (sec)</th></tr>"
            for lp in log_param_rows:
                html += (
                    f"<tr><td>{escape(str(lp.get('tagPath')))}</td>"
                    f"<td>{escape(str(lp.get('description')))}</td>"
                    f"<td>{escape(str(lp.get('rangeMin')))}</td>"
                    f"<td>{escape(str(lp.get('rangeMax')))}</td>"
                    f"<td>{escape(str(lp.get('logInterval')))}</td></tr>"
                )
            html += "</table>"

        filename = f"node_{safe}_{stamp}.pdf"
        path = d / filename
        build_pdf_html(title, html, path)
        return file_payload(
            path,
            filename,
            "application/pdf",
            title,
            {
                "node_id": nid,
                "node_name": node_name_s,
                "tag_count": len(tag_rows),
                "alarm_count": len(alarm_rows),
                "log_params": len(log_param_rows),
            },
            cfg=cfg,
        )

    sections: list[dict[str, Any]] = [
        {
            "heading": "Node Info",
            "text": f"ID: {nid} | Type: {node.get('nType')} | Path: {node.get('nPath')}",
        }
    ]
    if tag_rows:
        sections.append(
            {
                "heading": "Live Tag Values",
                "table": {
                    "headers": ["Tag", "Value", "Last Read"],
                    "rows": [[t.get("tagPath"), t.get("tagValue"), t.get("readTime")] for t in tag_rows],
                },
            }
        )
    if alarm_rows:
        sections.append(
            {
                "heading": "Alarm Parameters",
                "table": {
                    "headers": ["Parameter", "Description", "Status"],
                    "rows": [
                        [
                            a.get("tagPath"),
                            a.get("description") or "",
                            "ACTIVE" if int(a.get("alarmStateCurrent") or 0) > 0 else "Normal",
                        ]
                        for a in alarm_rows
                    ],
                },
            }
        )
    if log_param_rows:
        sections.append(
            {
                "heading": "Log Parameters",
                "table": {
                    "headers": ["Tag", "Description", "Min", "Max", "Interval(sec)"],
                    "rows": [
                        [
                            lp.get("tagPath"),
                            lp.get("description") or "",
                            lp.get("rangeMin"),
                            lp.get("rangeMax"),
                            lp.get("logInterval"),
                        ]
                        for lp in log_param_rows
                    ],
                },
            }
        )

    filename = f"node_{safe}_{stamp}.docx"
    path = d / filename
    build_docx_report(title, sections, path)
    return file_payload(
        path,
        filename,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        title,
        {
            "node_id": nid,
            "node_name": node_name_s,
            "tag_count": len(tag_rows),
            "alarm_count": len(alarm_rows),
            "log_params": len(log_param_rows),
        },
        cfg=cfg,
    )
